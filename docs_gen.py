"""Geração dos documentos oficiais em Word a partir dos dados da solicitação.

Cada função devolve um BytesIO com um .docx pronto para download.
Os modelos seguem os anexos da IN GAR/RET/UFF 058/2023.
"""
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


AZUL = RGBColor(0x1F, 0x4E, 0x78)


def _data_br(iso):
    if not iso:
        return ""
    try:
        return datetime.strptime(iso, "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return iso


def _base(titulo):
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Arial"
    normal.size = Pt(10)
    cab = doc.add_paragraph()
    cab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, tam in (("Serviço Público Federal", 9), ("Ministério da Educação", 9),
                     ("Universidade Federal Fluminense", 9)):
        r = cab.add_run(txt + "\n"); r.font.size = Pt(tam); r.font.color.rgb = AZUL
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rh = h.add_run(titulo); rh.bold = True; rh.font.size = Pt(12); rh.font.color.rgb = AZUL
    return doc


def _secao(doc, titulo):
    p = doc.add_paragraph()
    r = p.add_run(titulo); r.bold = True; r.font.color.rgb = AZUL
    return p


def _tabela(doc, linhas):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for rotulo, valor in linhas:
        c = t.add_row().cells
        c[0].text = rotulo
        c[1].text = valor or ""
        for par in c[0].paragraphs:
            for run in par.runs:
                run.bold = True
    return t


def _assinaturas(doc, esquerda, direita):
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for cel, txt in zip(t.rows[0].cells, (esquerda, direita)):
        cel.text = ""
        p = cel.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n\n_______________________________\n").font.size = Pt(9)
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(9)


def _salvar(doc):
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------- Anexo II - Requisição ----------------
def anexo_ii(perfil, resp):
    doc = _base("ANEXO II - REQUISIÇÃO DE DIÁRIAS E PASSAGENS E AFASTAMENTOS NO SCDP")
    doc.add_paragraph("Data da Solicitação: " + datetime.today().strftime("%d/%m/%Y"))

    _secao(doc, "1. PROPONENTE")
    _tabela(doc, [
        ("Proponente / Concedente", resp.get("chefe_departamento", "")),
        ("Unidade Responsável", perfil.get("departamento", "")),
    ])

    _secao(doc, "3. DADOS PESSOAIS (extraídos do SIAPE)")
    _tabela(doc, [
        ("Proposto", perfil.get("nome", "")),
        ("SIAPE", perfil.get("siape", "")),
        ("Cargo / Função", perfil.get("cargo", "")),
        ("Lotação / Órgão", perfil.get("departamento", "")),
        ("Telefone", perfil.get("telefone", "")),
        ("E-mail", perfil.get("email", "")),
    ])

    _secao(doc, "5. DADOS DA VIAGEM")
    _tabela(doc, [
        ("Passagens", resp.get("precisa_passagem", "")),
        ("Modo da passagem", resp.get("modo_passagem", "")),
        ("Diárias", resp.get("precisa_diaria", "")),
        ("Adicional de deslocamento", resp.get("adicional_deslocamento", "")),
        ("Bagagem despachada", resp.get("bagagem_despachada", "")),
        ("Possui seguro internacional", resp.get("apolice") and "Sim" or ""),
    ])

    _secao(doc, "TRECHOS / PERCURSOS")
    _tabela(doc, [
        ("Ida", resp.get("trecho_ida", "")),
        ("Volta", resp.get("trecho_volta", "")),
    ])

    _secao(doc, "PARTICIPAÇÃO DO PROPOSTO NO EVENTO / MISSÃO")
    _tabela(doc, [
        ("Data e hora de início da missão", resp.get("inicio_missao", "") or _data_br(resp.get("data_inicio", ""))),
        ("Data e hora do fim da missão", resp.get("fim_missao", "") or _data_br(resp.get("data_fim", ""))),
    ])

    _secao(doc, "6. OBJETIVO DA VIAGEM")
    _tabela(doc, [
        ("Motivo", resp.get("motivo", "")),
        ("Destino", (resp.get("destino", "") or resp.get("cidade_destino", "")) +
                    ((" - " + resp.get("uf_destino")) if resp.get("uf_destino") else "") +
                    ((" - " + resp.get("pais_destino")) if resp.get("pais_destino") else "")),
        ("Descrição do motivo, com vínculo às atividades da UFF", resp.get("descricao_motivo", "")),
        ("Nº do processo no SEI", resp.get("num_sei", "")),
    ])

    doc.add_paragraph()
    nota = doc.add_paragraph()
    rn = nota.add_run("Este formulário somente será considerado válido após assinatura do servidor e do chefe da unidade.")
    rn.italic = True; rn.font.size = Pt(9)
    _assinaturas(doc, "Assinatura do proposto\n" + perfil.get("nome", ""),
                 "Assinatura da chefia / proponente\n" + resp.get("chefe_departamento", ""))
    return _salvar(doc)


# ---------------- Anexo V - Termo de Renúncia ----------------
def termo_renuncia(perfil, resp):
    doc = _base("ANEXO V - TERMO DE RENÚNCIA")
    aplica = resp.get("renuncia_aplica", "")
    objeto = {
        "Renuncio a diárias": "diárias",
        "Renuncio a passagens": "passagens",
        "Renuncio a diárias e passagens": "diárias e passagens",
    }.get(aplica, "diárias e/ou passagens")

    p = doc.add_paragraph()
    p.add_run(
        "Pelo presente termo, Eu, " + perfil.get("nome", "") +
        ", matrícula SIAPE nº " + (perfil.get("siape") or "____") +
        ", " + (perfil.get("cargo") or "____") +
        ", lotado(a) no(a) " + (perfil.get("departamento") or "____") +
        ", li e concordo com os termos da IN GAR/RET/UFF 058/2023 e RENUNCIO EXPRESSAMENTE ao recebimento de " +
        objeto + ", a que eventualmente eu faça jus, destinado(a) à participação no evento " +
        (resp.get("descricao_motivo") or "____") +
        ", a ser realizado na cidade de " +
        (resp.get("destino") or resp.get("cidade_destino") or "____") +
        ", no período de " + _data_br(resp.get("data_inicio", "")) + " a " + _data_br(resp.get("data_fim", "")) +
        ", pelo motivo: " + (resp.get("renuncia_motivo") or "____") + "."
    )
    doc.add_paragraph()
    doc.add_paragraph("Niterói, " + datetime.today().strftime("%d/%m/%Y") + ".")
    _assinaturas(doc, "Assinatura do PROPOSTO / SERVIDOR\n" + perfil.get("nome", ""), " ")
    return _salvar(doc)


# ---------------- Anexo III / IV - Relatório de viagem ----------------
def relatorio(perfil, resp, internacional=False):
    titulo = ("ANEXO IV - RELATÓRIO DE VIAGEM INTERNACIONAL NO SCDP" if internacional
              else "ANEXO III - RELATÓRIO DE VIAGEM NACIONAL NO SCDP")
    doc = _base(titulo)
    _secao(doc, "Identificação")
    _tabela(doc, [
        ("Nome do proposto", perfil.get("nome", "")),
        ("Órgão / Setor", perfil.get("departamento", "")),
        ("Nº da PCDP", ""),
        ("Nº do processo no SEI", resp.get("num_sei", "")),
        ("Data de início da viagem", _data_br(resp.get("data_inicio", ""))),
        ("Data fim da viagem", _data_br(resp.get("data_fim", ""))),
        ("Roteiro da viagem", resp.get("trecho_ida", "")),
    ])
    _secao(doc, "1. Relatório de viagem resumido (preencher após a viagem)")
    doc.add_paragraph("Informar o tema do evento, evidenciando se o objetivo da viagem foi atingido.")
    doc.add_paragraph("_" * 90)
    doc.add_paragraph("_" * 90)
    if internacional:
        _secao(doc, "Conclusões alcançadas e benefícios para a Educação")
        doc.add_paragraph("_" * 90)
    _secao(doc, "2. A viagem foi realizada dentro do período previsto?")
    doc.add_paragraph("(  ) Sim    (  ) Não. Justifique.")
    _secao(doc, "3. Justificativa para prestação de contas fora do prazo (5 dias do término)")
    doc.add_paragraph("_" * 90)
    _assinaturas(doc, "Assinatura do proposto\n" + perfil.get("nome", ""),
                 "Assinatura da chefia / proponente")
    return _salvar(doc)


# ---------------- montagem por situação ----------------
def documentos_da_situacao(situacao, perfil, resp):
    """Retorna lista de (rotulo, nome_arquivo, tipo) disponíveis para a situação."""
    docs = [("Anexo II - Requisição", "anexo_ii.docx", "anexo_ii")]
    if resp.get("renuncia_aplica") and resp.get("renuncia_aplica") != "Não se aplica":
        docs.append(("Anexo V - Termo de Renúncia", "termo_renuncia.docx", "termo_renuncia"))
    if situacao == "internacional":
        docs.append(("Anexo IV - Relatório internacional", "relatorio_internacional.docx", "relatorio_int"))
    else:
        docs.append(("Anexo III - Relatório nacional", "relatorio_nacional.docx", "relatorio_nac"))
    return docs


def gerar(tipo, perfil, resp):
    if tipo == "anexo_ii":
        return anexo_ii(perfil, resp), "anexo_ii.docx"
    if tipo == "termo_renuncia":
        return termo_renuncia(perfil, resp), "termo_renuncia.docx"
    if tipo == "relatorio_int":
        return relatorio(perfil, resp, internacional=True), "relatorio_internacional.docx"
    if tipo == "relatorio_nac":
        return relatorio(perfil, resp, internacional=False), "relatorio_nacional.docx"
    return None, None
